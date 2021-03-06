import datetime
from selenium import webdriver
from selenium.common.exceptions import TimeoutException,NoSuchElementException
from selenium.webdriver.common.keys import Keys  
from docx.shared import Inches
import docx,os
# from selenium.webdriver.chrome.options import Options
import time,sys,os
from sys import platform
logFile = open("log.txt","a+")
logFile.write("\nStarted at: " + str(datetime.datetime.now()))
cur_path = sys.path[0]
def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.dirname(__file__)
    return os.path.join(base_path, relative_path)


if platform == "linux" or platform == "linux2":
    # linux
    path = resource_path('driver/chromedriver')
else:
    path = resource_path('driver/chromedriver.exe')
# chrome_options = Options()
# chrome_options.add_argument("--disable-extensions")
# chrome_options.add_argument("--headless")
# chrome_options.add_argument("--disable-gpu")
# chrome_options.add_argument("--window-size=1920,1080")
# chrome_options.add_argument("--no-sandbox") # linux only
# chrome_options.headless = True # also works
# driver = webdriver.Chrome()
    # Windows...
print("\n\nProcessing.....")

# driver =webdriver.Chrome(path,options=chrome_options)
driver =webdriver.Chrome(path)

            
driver.maximize_window()
# open link
# driver.set_page_load_timeout(120)
driver.set_page_load_timeout(30)

# data = pd.read_csv("Equity.csv")
# data = data['Security Id'] 
# shortcode = "HDFC"

try:
    driver.get("https://www.restojobs.ca/en/applicants")
    logFile.write("\nsuccess : Loaded...")
    print("success : Loaded...")
except TimeoutException as e:
    logFile.write("\ninfo : website taking too long to load...stopped")
    print("info : website taking too long to load...stopped")
    # driver.refresh()
if(os.path.exists("images")):
    print("Folder is already there")
else:
    os.mkdir("images")
    print("creating new folder")
try:    
    search = driver.find_element_by_name("place")
    search.send_keys("Montreal")
    time.sleep(1)
    search.send_keys(Keys.DOWN)
    search.send_keys(Keys.ENTER)
except Exception as e:
    print(e)
time.sleep(5)
try:    
    driver.find_element_by_partial_link_text("Cook & Chef").click()
except Exception as e:
    print(e)
time.sleep(5)

profile_links = []
try:
    print("info : Finding links...")
    driver.execute_script("window.scrollTo(0,document.body.scrollHeight);")
    for i in range(2):
        driver.execute_script("window.scrollTo(0,document.body.scrollHeight);")
        driver.find_element_by_css_selector("a.load-more").click()
        time.sleep(3)
    links = driver.find_elements_by_css_selector("a.background-link")
    for link in links:
        print(link.get_attribute("href"))
        profile_links.append(link.get_attribute("href"))
except Exception as e:
    print(e)

try:
    document = docx.Document()
    print("info : Extracting data...")
    for link in profile_links:
        driver.get(link)
        driver.execute_script("window.scrollTo(0,340);")
        txt = driver.find_element_by_class_name("applicant-info").text
        # print(txt)
        filepath = "images/img.png"

        driver.save_screenshot(filepath)

        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(filepath,width=Inches(6.50))
        txt = txt.split("\n")
        for data in txt:
            r.add_text(data)
            r.add_break()
    document.save('data.docx')
except Exception as e:
    print(e)
logFile.write("\nsuccess : complete")
print("success : complete")

logFile.close()